import re
import io
from typing import Optional
import pdfplumber
from docx import Document

from app.schemas.cv_data import CvData


class CvParser:
    """Tool for parsing CV files and extracting text and basic information"""
    
    def __init__(self):
        # Common email regex
        self.email_pattern = re.compile(
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        )
        # Phone patterns (various formats)
        self.phone_patterns = [
            re.compile(r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}'),
            re.compile(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'),
            re.compile(r'\+\d{1,3}\s?\d{6,12}'),
        ]
    
    def parse(self, file_content: bytes, filename: str) -> CvData:
        """Parse CV file and extract data"""
        ext = filename.lower().split('.')[-1]
        
        if ext == 'pdf':
            raw_text = self._extract_pdf_text(file_content)
        elif ext in ['doc', 'docx']:
            raw_text = self._extract_docx_text(file_content)
        else:
            raw_text = file_content.decode('utf-8', errors='ignore')
        
        # Extract basic info using regex
        email = self._extract_email(raw_text)
        phone = self._extract_phone(raw_text)
        first_name, last_name = self._extract_name(raw_text)
        
        return CvData(
            raw_text=raw_text,
            first_name=first_name,
            last_name=last_name,
            email=email,
            phone=phone
        )
    
    def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF using pdfplumber"""
        text_parts = []
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception as e:
            # Fallback or log error
            print(f"PDF parsing error: {e}")
        return '\n'.join(text_parts)
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        text_parts = []
        try:
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
        except Exception as e:
            print(f"DOCX parsing error: {e}")
        return '\n'.join(text_parts)
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract first email address from text"""
        match = self.email_pattern.search(text)
        return match.group(0) if match else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number from text"""
        for pattern in self.phone_patterns:
            match = pattern.search(text)
            if match:
                phone = match.group(0)
                # Clean up and validate
                digits = re.sub(r'\D', '', phone)
                if 7 <= len(digits) <= 15:
                    return phone
        return None
    
    def _extract_name(self, text: str) -> tuple[Optional[str], Optional[str]]:
        """
        Extract name from CV - typically at the top.
        This is a heuristic approach, LLM will refine this.
        """
        lines = text.strip().split('\n')
        
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            
            # Skip lines that look like headers or contact info
            if not line or '@' in line or any(c.isdigit() for c in line[:3]):
                continue
            
            # Skip common headers
            skip_words = ['resume', 'cv', 'curriculum', 'vitae', 'profile', 'summary', 
                         'objective', 'experience', 'education', 'skills', 'contact']
            if any(word in line.lower() for word in skip_words):
                continue
            
            # Try to parse as name (2-4 words, all alphabetic)
            words = line.split()
            if 2 <= len(words) <= 4:
                if all(word.replace('-', '').replace("'", '').isalpha() for word in words):
                    # Likely a name
                    return words[0], ' '.join(words[1:])
        
        return None, None


# Singleton instance
cv_parser = CvParser()
